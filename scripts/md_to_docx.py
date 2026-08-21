#!/usr/bin/env python3
"""Minimal Markdown -> .docx converter for the project report.

Handles: ATX headings (#..######), bullet lists (- ), GitHub tables,
horizontal rules (---), bold (**x**), inline code (`x`), and paragraphs.
Good enough for docs/PROJECT_REPORT.md; not a general Markdown engine.
"""
import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_runs(paragraph, text):
    """Render **bold** and `code` inline spans into a paragraph."""
    token = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    for part in token.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.color.rgb = RGBColor(0x8A, 0x2B, 0xE2)
        else:
            paragraph.add_run(part)


def parse_table(lines, i, doc):
    header = [c.strip() for c in lines[i].strip().strip("|").split("|")]
    rows = []
    j = i + 2  # skip the |---|---| separator
    while j < len(lines) and lines[j].lstrip().startswith("|"):
        rows.append([c.strip() for c in lines[j].strip().strip("|").split("|")])
        j += 1
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for k, cell in enumerate(header):
        p = table.rows[0].cells[k].paragraphs[0]
        add_runs(p, cell)
        for run in p.runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for k in range(len(header)):
            val = row[k] if k < len(row) else ""
            add_runs(cells[k].paragraphs[0], val)
    doc.add_paragraph()
    return j


def main(md_path, docx_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph().add_run("_" * 40).font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            heading = doc.add_heading(level=min(level, 4))
            add_runs(heading, m.group(2))
            i += 1
            continue

        # table
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|?$", lines[i + 1].strip()):
            i = parse_table(lines, i, doc)
            continue

        # bullet
        mb = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if mb:
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, mb.group(2))
            i += 1
            continue

        # numbered
        mn = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if mn:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, mn.group(1))
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
